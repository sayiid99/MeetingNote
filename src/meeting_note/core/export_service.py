from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from meeting_note.core.contracts import TranscriptDocument, TranscriptSegment, TranslationDocument
from meeting_note.core.formatting import format_transcript_document


class ExportService:
    def export_text(self, content: str, output_path: Path) -> Path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(content, encoding="utf-8")
        return output_path

    def export_markdown(self, title: str, content: str, output_path: Path) -> Path:
        markdown = f"# {title}\n\n{content.strip()}\n"
        return self.export_text(markdown, output_path)

    def export_docx(self, title: str, content: str, output_path: Path) -> Path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        document.add_heading(title, level=1)
        for paragraph in content.splitlines():
            if not paragraph.strip():
                continue
            run = document.add_paragraph().add_run(paragraph)
            run.font.size = Pt(11)
        document.save(output_path)
        return output_path

    def export_pdf(self, title: str, content: str, output_path: Path) -> Path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        styles = getSampleStyleSheet()
        body_style = ParagraphStyle(
            "MeetingNoteBody",
            parent=styles["Normal"],
            fontSize=11,
            leading=16,
        )
        title_style = ParagraphStyle(
            "MeetingNoteTitle",
            parent=styles["Title"],
            fontSize=18,
            leading=22,
        )
        story = [Paragraph(self._escape_pdf_text(title), title_style), Spacer(1, 12)]
        for paragraph in content.splitlines():
            if paragraph.strip():
                story.append(Paragraph(self._escape_pdf_text(paragraph), body_style))
                story.append(Spacer(1, 6))
        SimpleDocTemplate(str(output_path), pagesize=A4).build(story)
        return output_path

    def export_transcript_txt(self, transcript: TranscriptDocument, output_path: Path) -> Path:
        return self.export_text(format_transcript_document(transcript), output_path)

    def export_translation_txt(self, translation: TranslationDocument, output_path: Path) -> Path:
        return self.export_text(translation.translated_text, output_path)

    def export_bilingual_markdown(self, translation: TranslationDocument, output_path: Path) -> Path:
        return self.export_text(translation.bilingual_text or translation.translated_text, output_path)

    def export_srt(self, transcript: TranscriptDocument, output_path: Path, show_speakers: bool = True) -> Path:
        normalized_segments = self._normalized_srt_segments(transcript.segments)
        blocks = []
        for index, segment in enumerate(normalized_segments, start=1):
            blocks.append(self._format_srt_block(index, segment, show_speakers))
        return self.export_text("\n\n".join(blocks), output_path)

    @classmethod
    def _format_srt_block(cls, index: int, segment: TranscriptSegment, show_speakers: bool) -> str:
        text = segment.text.replace("\n", " ").replace("\r", " ")
        if show_speakers and segment.speaker_id:
            text = f"[{segment.speaker_id}] {text}"
        return "\n".join(
            [
                str(index),
                f"{cls.format_srt_timestamp(segment.start_time)} --> {cls.format_srt_timestamp(segment.end_time)}",
                text,
            ]
        )

    @staticmethod
    def format_srt_timestamp(seconds: float) -> str:
        safe_ms = max(0, int(round(seconds * 1000)))
        milliseconds = safe_ms % 1000
        total_seconds = safe_ms // 1000
        seconds_part = total_seconds % 60
        total_minutes = total_seconds // 60
        minutes = total_minutes % 60
        hours = total_minutes // 60
        return f"{hours:02d}:{minutes:02d}:{seconds_part:02d},{milliseconds:03d}"

    @staticmethod
    def _escape_pdf_text(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    @classmethod
    def _normalized_srt_segments(cls, segments: list[TranscriptSegment]) -> list[TranscriptSegment]:
        normalized: list[TranscriptSegment] = []
        cursor = 0.0
        for segment in segments:
            start_time = cls._safe_float(segment.start_time)
            end_time = cls._safe_float(segment.end_time)

            if start_time < cursor:
                start_time = cursor
            if start_time <= 0 and end_time <= 0:
                start_time = cursor
            if end_time <= start_time:
                end_time = start_time + cls._estimate_segment_duration(segment.text)
            if end_time <= start_time:
                end_time = start_time + 1.0

            cursor = end_time
            normalized.append(
                TranscriptSegment(
                    id=segment.id,
                    text=segment.text,
                    start_time=start_time,
                    end_time=end_time,
                    speaker_id=segment.speaker_id,
                    is_edited=segment.is_edited,
                )
            )
        return normalized

    @staticmethod
    def _safe_float(value: float | int | str | None) -> float:
        try:
            return max(0.0, float(value or 0.0))
        except (TypeError, ValueError):
            return 0.0

    @staticmethod
    def _estimate_segment_duration(text: str) -> float:
        compact = " ".join(text.split())
        if not compact:
            return 1.5

        words = re.findall(r"[A-Za-z0-9']+", compact)
        if words:
            estimated = len(words) / 2.8
        else:
            estimated = len(compact) / 4.2
        return max(1.2, min(300.0, estimated))
